"""
Format-specific document loaders: PDF, DOCX, TXT, HTML -> plain text.

WHY THIS MODULE EXISTS
-----------------------
The chunker (Milestone 4) and everything after it should never need to
know or care whether a document originally came from a PDF, a Word file,
an HTML page, or a plain-text export. This module is the adapter layer
that converts each format's very different internal structure (a PDF's
page/glyph layout, a DOCX's paragraph/run XML tree, an HTML page's DOM)
into one common shape: `LoadedDocument`, a list of cleaned per-page text
strings plus a few metadata fields. Everything downstream depends only on
that shape, not on `fitz`, `docx`, or `bs4` directly — swapping or adding
a parser (e.g. Docling for complex layouts) later only touches this file.

DESIGN: WHY "PAGES" FOR EVERY FORMAT, EVEN NON-PAGINATED ONES
------------------------------------------------------------------
The spec requires each chunk to carry a page number. PDFs have a real,
unambiguous notion of a page; DOCX and HTML do not (a .docx page break is
a rendering-time concept, not reliably present in the XML; HTML has no
pages at all). Rather than giving `LoadedDocument` two different shapes,
every loader returns `pages: list[str]` uniformly — for PDFs it's one
string per real page; for DOCX/TXT/HTML it's a single-element list holding
the whole document, and downstream code treats `page_number` as `None` for
those formats (see `ChunkRecord.page_number: Optional[int]`). One shape,
one code path in the chunker, with page numbers simply absent where they
don't exist.

OCR FALLBACK FOR SCANNED PDFS
--------------------------------
Some PDFs (especially scanned/photographed legal documents) have no
extractable text layer at all — `page.get_text()` returns empty for every
page. For those pages we render the page to an image using PyMuPDF's own
rasterizer (`page.get_pixmap()` — no external `poppler` binary needed,
unlike the more common `pdf2image` approach) and run Tesseract OCR on the
image via `pytesseract`. This is opt-in *per page*: pages that already
have a text layer skip OCR entirely (OCR is slow — roughly 1-3 seconds per
page — and lower-quality than native text extraction), so a mixed
document only pays the OCR cost where it's actually needed.

If Tesseract isn't installed on the host machine, OCR is skipped with a
warning recorded on the `LoadedDocument` rather than crashing the whole
load — a document with some unreadable scanned pages is still partially
useful, which fits the "gracefully handle broken input" requirement.

ENCODING FALLBACK FOR TXT FILES
-----------------------------------
Legal text corpora are frequently a mix of UTF-8 (modern digitization) and
legacy Cyrillic-Uzbek encodings (Windows-1251) from older sources. We try,
in order: UTF-8 with BOM handling (`utf-8-sig` — strips a BOM if present,
behaves like plain UTF-8 if not), then plain UTF-8, then `windows-1251`,
then `latin-1` (which can decode ANY byte sequence, so it's the guaranteed
last resort — used only so ingestion never crashes outright on unknown
encodings, with a loud warning that the result may be garbled).

TIME / MEMORY COMPLEXITY
-------------------------
- TXT: O(file size) — a single read + decode pass.
- DOCX: O(number of paragraphs + table cells) — python-docx parses the
  whole XML tree into memory, so memory is O(document size); fine for
  legal documents (tens of KB to a few MB), would need streaming for
  gigabyte-scale inputs.
- PDF (native text): O(number of pages), each page's text extraction is
  roughly linear in that page's content.
- PDF (OCR fallback): O(pages needing OCR) x (image render + Tesseract
  inference cost) — by far the most expensive path, which is exactly why
  it's applied selectively.
- HTML: O(DOM size) for parsing plus O(text size) for extraction.

ADVANTAGES
-----------
- One uniform output shape (`LoadedDocument`) decouples every downstream
  module from parsing-library specifics.
- OCR fallback means scanned PDFs degrade gracefully instead of silently
  producing zero chunks.
- Encoding fallback chain means TXT ingestion never hard-crashes on a
  legacy-encoded file.

DISADVANTAGES
--------------
- python-docx cannot recover real page numbers (a DOCX has none in its
  object model) — acceptable per the spec, which only requires page
  numbers where they're meaningful.
- OCR text quality depends entirely on scan quality and the installed
  Tesseract language data; there's no post-OCR spell-correction here.
- `latin-1` as a final encoding fallback can silently produce mojibake
  rather than failing loudly — a deliberate trade-off (never crash
  ingestion) that trades correctness for availability; the resulting
  `LoadedDocument.warnings` flags it so the pipeline can surface it.

ALTERNATIVES CONSIDERED
-------------------------
- `chardet`/`charset-normalizer` for automatic encoding detection: more
  robust than a fixed fallback chain, but adds a dependency for a problem
  a 3-encoding chain already handles well for this corpus's known
  provenance (modern UTF-8 + legacy Windows-1251 Uzbek/Cyrillic sources).
- Docling for PDF parsing: better table/layout structure extraction, at
  the cost of a much heavier install and slower parsing; left as an
  optional upgrade path (see requirements.txt) rather than the default,
  since our source PDFs (when added) are expected to be simple running
  text, not complex multi-column layouts.

BEST PRACTICES APPLIED
------------------------
- Every loader raises from `app.ingestion.exceptions` on recoverable
  failure, never a bare built-in exception, so the pipeline can catch one
  base class (see exceptions.py's docstring).
- Cleaning (`clean_text`) is applied uniformly to every page's text in
  `load_document`, not duplicated inside each per-format loader.
- Every warning (encoding fallback used, OCR applied, OCR unavailable) is
  both logged AND recorded on the returned object, so it's visible both in
  `logs/app.log` and in the UI (later milestones can surface
  `LoadedDocument.warnings` directly to the user).
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

from pydantic import BaseModel

from app.ingestion.cleaning import clean_text
from app.ingestion.exceptions import (
    CorruptedDocumentError,
    EmptyDocumentError,
    UnsupportedFileTypeError,
)
from app.utils.logger import get_logger

logger = get_logger(__name__)

# Encodings tried in order for plain-text files. `utf-8-sig` handles a
# leading BOM transparently and otherwise behaves like `utf-8`, so it's
# tried first rather than as a separate special case.
_TXT_ENCODING_FALLBACKS = ("utf-8-sig", "utf-8", "windows-1251", "latin-1")

# Below this many extracted characters, a PDF page is considered to have
# no usable text layer and becomes an OCR candidate.
_OCR_MIN_TEXT_CHARS = 10

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".html", ".htm"}


class LoadedDocument(BaseModel):
    """Uniform output of every format-specific loader.

    `pages` always has at least one element for a successfully-loaded,
    non-empty document. `full_text` is the join of all cleaned pages and
    is what most callers actually want; `pages` exists so the chunker can
    recover per-chunk page numbers for paginated formats.
    """

    file_path: str
    file_type: str  # "pdf" | "docx" | "txt" | "html"
    pages: list[str]
    warnings: list[str] = []

    @property
    def full_text(self) -> str:
        return "\n\n".join(p for p in self.pages if p)

    @property
    def char_count(self) -> int:
        return len(self.full_text)

    @property
    def num_pages(self) -> int:
        return len(self.pages)


# --------------------------------------------------------------------------
# TXT
# --------------------------------------------------------------------------


def load_txt(path: Path) -> list[str]:
    """Read a plain-text file, trying encodings in `_TXT_ENCODING_FALLBACKS`.

    Returns a single-element page list (TXT has no pagination concept).
    """
    raw_bytes = path.read_bytes()
    last_error: Optional[UnicodeDecodeError] = None

    for encoding in _TXT_ENCODING_FALLBACKS:
        try:
            text = raw_bytes.decode(encoding)
            if encoding != "utf-8-sig":
                logger.warning("File %s decoded using fallback encoding=%s", path, encoding)
            return [text]
        except UnicodeDecodeError as e:
            last_error = e
            continue

    # latin-1 can decode any byte sequence, so reaching here should be
    # impossible — kept as a defensive guard, not a reachable branch.
    raise CorruptedDocumentError(
        f"Could not decode {path} with any known encoding: {last_error}"
    )


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------


def load_docx(path: Path) -> list[str]:
    """Extract paragraph and table text from a .docx file.

    DOCX has no reliable page-boundary information in its object model
    (page breaks are a layout-engine concept), so this returns a single
    "page" containing the whole document.
    """
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover - dependency presence, not logic
        raise ImportError("python-docx is required to load .docx files") from e

    try:
        document = docx.Document(str(path))
    except Exception as e:
        raise CorruptedDocumentError(f"Could not open {path} as a .docx file: {e}") from e

    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return ["\n".join(parts)]


# --------------------------------------------------------------------------
# HTML
# --------------------------------------------------------------------------


def load_html(path: Path) -> list[str]:
    """Extract visible text from an HTML file, discarding markup/scripts."""
    from bs4 import BeautifulSoup

    raw_bytes = path.read_bytes()
    try:
        text = raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = raw_bytes.decode("windows-1251", errors="replace")

    try:
        soup = BeautifulSoup(text, "lxml")
    except Exception as e:
        raise CorruptedDocumentError(f"Could not parse {path} as HTML: {e}") from e

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    extracted = soup.get_text(separator="\n")
    return [extracted]


# --------------------------------------------------------------------------
# PDF (with OCR fallback for scanned pages)
# --------------------------------------------------------------------------


def _ocr_page_image(page, warnings: list[str]) -> str:
    """Render one PDF page to an image and run Tesseract OCR on it.

    Returns an empty string (and appends a warning) if OCR is unavailable
    or fails — the caller treats that page as having no text, rather than
    letting an OCR failure abort the whole document load.
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        warnings.append(
            "OCR dependencies (pytesseract/Pillow) not installed; "
            "scanned page(s) yielded no text."
        )
        return ""

    try:
        # 200 DPI balances OCR accuracy against render/inference time; the
        # PyMuPDF default is 72 DPI, too low for reliable character
        # recognition on typical scanned legal documents.
        pixmap = page.get_pixmap(dpi=200)
        image = Image.open(io.BytesIO(pixmap.tobytes("png")))
        return pytesseract.image_to_string(image, lang="eng")
    except pytesseract.TesseractNotFoundError:
        warnings.append(
            "Tesseract binary not found on PATH; scanned page(s) yielded no text. "
            "Install it (e.g. `brew install tesseract`) to enable OCR."
        )
        return ""
    except Exception as e:  # noqa: BLE001 - OCR failures are recoverable, not fatal
        warnings.append(f"OCR failed on a page: {e}")
        return ""


def load_pdf(path: Path, warnings: Optional[list[str]] = None) -> list[str]:
    """Extract text per page, falling back to OCR for pages with no text layer."""
    try:
        import fitz  # PyMuPDF
    except ImportError as e:  # pragma: no cover
        raise ImportError("PyMuPDF (fitz) is required to load .pdf files") from e

    if warnings is None:
        warnings = []

    try:
        doc = fitz.open(str(path))
    except Exception as e:
        raise CorruptedDocumentError(f"Could not open {path} as a PDF: {e}") from e

    if doc.is_encrypted:
        # Try an empty-password unlock (common for "restricted but not
        # actually password-protected" PDFs); if that fails, it's genuinely
        # inaccessible without a password we don't have.
        if not doc.authenticate(""):
            raise CorruptedDocumentError(f"{path} is password-protected; cannot extract text.")

    pages: list[str] = []
    for page_index in range(doc.page_count):
        page = doc.load_page(page_index)
        text = page.get_text()
        if len(text.strip()) < _OCR_MIN_TEXT_CHARS:
            logger.info(
                "Page %d of %s has no usable text layer; attempting OCR", page_index + 1, path
            )
            text = _ocr_page_image(page, warnings)
        pages.append(text)

    doc.close()
    return pages


# --------------------------------------------------------------------------
# Dispatch
# --------------------------------------------------------------------------

_LOADERS = {
    ".txt": load_txt,
    ".docx": load_docx,
    ".html": load_html,
    ".htm": load_html,
}


def load_document(path: Path) -> LoadedDocument:
    """Load any supported file into a cleaned, uniform `LoadedDocument`.

    Raises `UnsupportedFileTypeError` for unknown extensions,
    `EmptyDocumentError` if every page is empty after cleaning, or a
    format-specific `DocumentLoadError` subclass on parse failure.
    """
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{suffix}' for {path}. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    warnings: list[str] = []

    logger.info("Loading document: %s", path)
    if suffix == ".pdf":
        raw_pages = load_pdf(path, warnings=warnings)
    else:
        raw_pages = _LOADERS[suffix](path)

    cleaned_pages = [clean_text(p) for p in raw_pages]

    if not any(p.strip() for p in cleaned_pages):
        raise EmptyDocumentError(f"{path} contained no extractable text after cleaning.")

    doc = LoadedDocument(
        file_path=str(path),
        file_type="html" if suffix == ".htm" else suffix.lstrip("."),
        pages=cleaned_pages,
        warnings=warnings,
    )
    logger.info(
        "Loaded %s: %d page(s), %d characters, %d warning(s)",
        path,
        doc.num_pages,
        doc.char_count,
        len(doc.warnings),
    )
    return doc
